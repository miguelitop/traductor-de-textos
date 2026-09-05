"""Tests para las funciones de docx_handler.py."""

from docx import Document

from traductor.docx_handler import (
    _tokens_intactos,
    _link_traducible,
    asignar_traducciones_links,
    mergear_parrafos,
    recolectar_links_traducibles,
    UnidadTraducible,
    HyperlinkInfo,
    NotaInfo,
)


def test_tokens_intactos_correctos():
    """Tokens esperados y encontrados coinciden: True."""
    unidad = UnidadTraducible(
        texto="dummy",
        parrafo="mock",
        traduccion="texto ⟦1⟧ con enlace y nota ⟦2⟧",
        hyperlinks={1: HyperlinkInfo(url="https://example.com", r_id="r1")},
        notas={2: NotaInfo(tipo="footnote", nota_id="1")},
    )
    assert _tokens_intactos(unidad) is True


def test_tokens_intactos_faltante():
    """Falta un token esperado en la traducción: False."""
    unidad = UnidadTraducible(
        texto="dummy",
        parrafo="mock",
        traduccion="texto ⟦1⟧ sin la nota",
        hyperlinks={1: HyperlinkInfo(url="https://example.com", r_id="r1")},
        notas={2: NotaInfo(tipo="footnote", nota_id="1")},
    )
    assert _tokens_intactos(unidad) is False


def test_tokens_intactos_sobrante():
    """Aparece un token no esperado en la traducción: False."""
    unidad = UnidadTraducible(
        texto="dummy",
        parrafo="mock",
        traduccion="texto ⟦1⟧ con ⟦2⟧ y extra ⟦3⟧",
        hyperlinks={1: HyperlinkInfo(url="https://example.com", r_id="r1")},
        notas={2: NotaInfo(tipo="footnote", nota_id="1")},
    )
    assert _tokens_intactos(unidad) is False


def test_tokens_intactos_sin_tokens():
    """Sin tokens esperados ni en la traducción: True."""
    unidad = UnidadTraducible(
        texto="dummy",
        parrafo="mock",
        traduccion="texto plano sin tokens",
    )
    assert _tokens_intactos(unidad) is True


def test_tokens_intactos_solo_hyperlinks():
    """Solo hyperlinks, sin notas: True si coinciden."""
    unidad = UnidadTraducible(
        texto="dummy",
        parrafo="mock",
        traduccion="enlace ⟦1⟧ y otro ⟦2⟧",
        hyperlinks={
            1: HyperlinkInfo(url="https://a.com", r_id="r1"),
            2: HyperlinkInfo(url="https://b.com", r_id="r2"),
        },
    )
    assert _tokens_intactos(unidad) is True


# --- Traducción del texto visible de los hipervínculos ---


def test_link_traducible_prosa():
    """Texto con palabras: se traduce."""
    assert _link_traducible("the official report") is True
    assert _link_traducible("Fig. 3") is True


def test_link_traducible_descarta_urls_y_numeros():
    """URLs, dominios, mails y numeración de citas no se traducen."""
    for texto in ("https://example.com/a", "www.example.com", "example.com/doc",
                  "info@example.com", "[12]", "3.2", "  "):
        assert _link_traducible(texto) is False, texto


def test_recolectar_links_traducibles_filtra():
    """Solo se recolectan los links con texto traducible."""
    unidad = UnidadTraducible(
        texto="⟦1⟧ y ⟦2⟧",
        parrafo="mock",
        hyperlinks={
            1: HyperlinkInfo(url="https://a.com", r_id="r1", texto="ver el informe"),
            2: HyperlinkInfo(url="https://b.com", r_id="r2", texto="www.b.com"),
        },
    )
    links = recolectar_links_traducibles([unidad])
    assert [l.texto for l in links] == ["ver el informe"]


def test_asignar_traducciones_links_ok():
    """Una traducción válida se guarda y es la que se reinserta."""
    link = HyperlinkInfo(url="https://a.com", r_id="r1", texto="official report")
    assert asignar_traducciones_links([link], ["informe oficial"]) == 0
    assert link.texto_final == "informe oficial"


def test_asignar_traducciones_links_descarta_sospechosas():
    """Salidas que no parecen texto de link dejan el original."""
    casos = [
        "",
        "linea uno\nlinea dos",
        "[ERROR DE TRADUCCIÓN EN CHUNK 3]\n\noficial",
        "informe oficial, que en este contexto se refiere al documento publicado",
    ]
    for traduccion in casos:
        link = HyperlinkInfo(url="https://a.com", r_id="r1", texto="official report")
        assert asignar_traducciones_links([link], [traduccion]) == 1, traduccion
        assert link.texto_final == "official report"


def test_mergear_parrafos_renumera_tokens_y_fusiona_mapas():
    """Al mergear, los tokens del segundo párrafo se corren y sus links se conservan."""
    doc = Document()
    p1, p2 = doc.add_paragraph("a"), doc.add_paragraph("b")
    u1 = UnidadTraducible(
        texto="según ⟦1⟧ el estudio",
        parrafo=p1,
        hyperlinks={1: HyperlinkInfo(url="https://a.com", r_id="r1", texto="Smith")},
    )
    u2 = UnidadTraducible(
        texto="citado en ⟦1⟧ y ⟦2⟧",
        parrafo=p2,
        hyperlinks={1: HyperlinkInfo(url="https://b.com", r_id="r2", texto="Jones")},
        notas={2: NotaInfo(tipo="footnote", nota_id="7")},
    )

    resultado = mergear_parrafos([u1, u2])

    assert len(resultado) == 1
    unido = resultado[0]
    assert unido.texto == "según ⟦1⟧ el estudio citado en ⟦2⟧ y ⟦3⟧"
    assert unido.hyperlinks[1].r_id == "r1"
    assert unido.hyperlinks[2].r_id == "r2"
    assert unido.notas[3].nota_id == "7"

    unido.traduccion = unido.texto
    assert _tokens_intactos(unido) is True
