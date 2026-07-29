import re
from copy import deepcopy
from dataclasses import dataclass, field
from pathlib import Path
from xml.etree.ElementTree import tostring

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph



@dataclass
class HyperlinkInfo:
    """Info de un hipervínculo dentro de un párrafo."""
    url: str
    r_id: str
    anchor: str = ""
    rPr: object = None
    texto: str = ""  # texto visible original del link (se reinserta tal cual)


@dataclass
class NotaInfo:
    """Info de una referencia a nota al pie o final."""
    tipo: str  # "footnote" o "endnote"
    nota_id: str
    rPr: object = None


@dataclass
class UnidadTraducible:
    """Representa un párrafo (o parte de uno) traducible con referencia a sus runs."""
    texto: str
    parrafo: Paragraph
    traduccion: str = ""
    hyperlinks: dict = field(default_factory=dict)
    notas: dict = field(default_factory=dict)
    base_rPr: object = None
    run_indices: list[int] | None = None  # None = todo el párrafo; lista = solo esos runs


@dataclass
class ImagenTraducible:
    """Representa una imagen inline cuyo texto se va a traducir."""
    parrafo: Paragraph              # párrafo que contiene la imagen
    run_element: object             # elemento w:r con el w:drawing de la imagen
    imagen_bytes: bytes
    ancho_emu: int                  # ancho de la imagen en EMU
    alineacion_jc: str              # 'left' | 'center' | 'right' (del párrafo padre)
    traduccion: str | None = None   # se completa después de la pasada de visión


def _parrafo_tiene_imagen(parrafo: Paragraph) -> bool:
    """Detecta si un párrafo contiene imágenes inline."""
    for run in parrafo.runs:
        if run._element.findall(qn("w:drawing")):
            return True
        if run._element.findall(qn("w:pict")):
            return True
    return False


def _resolver_imagen_blob(parrafo: Paragraph, run_element) -> tuple[bytes, int] | None:
    """Devuelve (bytes_imagen, ancho_emu) del primer drawing/pict del run, o None."""
    # Caso 1: w:drawing → wp:inline → a:graphic → a:graphicData → pic:pic → pic:blipFill → a:blip
    drawing = run_element.find(qn("w:drawing"))
    if drawing is not None:
        # Buscar el ancho de wp:extent
        ancho_emu = 0
        for tag in (qn("wp:extent"),):
            extent = drawing.find(f".//{tag}")
            if extent is not None:
                try:
                    ancho_emu = int(extent.get("cx", "0"))
                except ValueError:
                    ancho_emu = 0
                break
        # Buscar el a:blip con r:embed
        blip = drawing.find(f".//{qn('a:blip')}")
        if blip is not None:
            r_embed = blip.get(qn("r:embed"))
            if r_embed:
                try:
                    image_part = parrafo.part.related_parts[r_embed]
                    return image_part.blob, ancho_emu
                except (KeyError, AttributeError):
                    return None

    # Caso 2: w:pict (legacy VML) — más raro, lo skipeamos por ahora
    return None


def _alineacion_parrafo(parrafo: Paragraph) -> str:
    """Devuelve 'left' | 'center' | 'right' del párrafo (default 'center' para imágenes)."""
    pPr = parrafo._element.find(qn("w:pPr"))
    if pPr is not None:
        jc = pPr.find(qn("w:jc"))
        if jc is not None:
            val = jc.get(qn("w:val"))
            if val in ("left", "center", "right"):
                return val
    # Default razonable para imágenes: centradas
    return "center"


def _extraer_imagenes_de_parrafo(parrafo: Paragraph) -> list[ImagenTraducible]:
    """Devuelve las imágenes inline del párrafo como ImagenTraducible."""
    imagenes = []
    alineacion = _alineacion_parrafo(parrafo)
    for run in parrafo.runs:
        run_el = run._element
        if not (run_el.findall(qn("w:drawing")) or run_el.findall(qn("w:pict"))):
            continue
        info = _resolver_imagen_blob(parrafo, run_el)
        if info is None:
            continue
        imagen_bytes, ancho_emu = info
        if not imagen_bytes:
            continue
        imagenes.append(ImagenTraducible(
            parrafo=parrafo, run_element=run_el,
            imagen_bytes=imagen_bytes, ancho_emu=ancho_emu,
            alineacion_jc=alineacion,
        ))
    return imagenes


def _extraer_texto_con_links(parrafo: Paragraph) -> tuple[str, dict, dict, object]:
    """Extrae texto del párrafo con tokens opacos para hipervínculos y notas.

    Cada hipervínculo o nota se reemplaza por un token opaco ⟦N⟧ que el modelo
    solo tiene que arrastrar sin tocar (no traduce nada dentro). El texto visible
    del link NO se manda al modelo: se guarda en el HyperlinkInfo y se reinserta
    tal cual el original al reconstruir. Es mucho más robusto que embeber el texto
    en el marcador, porque los modelos locales rompen los delimitadores cuando
    tienen que traducir texto adentro.

    El contador es único para links y notas (comparten numeración) para que los
    índices nunca colisionen.

    Retorna (texto_con_tokens, {indice: HyperlinkInfo}, {indice: NotaInfo}, rPr_base).
    """
    elem = parrafo._element
    link_map = {}
    nota_map = {}
    contador = 0
    text_parts = []
    base_rPr = None

    for child in elem:
        if child.tag == qn("w:r"):
            # Verificar si el run contiene una referencia a nota
            fn_ref = child.find(qn("w:footnoteReference"))
            en_ref = child.find(qn("w:endnoteReference"))
            if fn_ref is not None or en_ref is not None:
                ref_elem = fn_ref if fn_ref is not None else en_ref
                tipo = "footnote" if fn_ref is not None else "endnote"
                nota_id = ref_elem.get(qn("w:id"), "")
                contador += 1
                nota_rPr = child.find(qn("w:rPr"))
                nota_map[contador] = NotaInfo(
                    tipo=tipo, nota_id=nota_id,
                    rPr=deepcopy(nota_rPr) if nota_rPr is not None else None,
                )
                text_parts.append(_token(contador))
            else:
                for t_elem in child.findall(qn("w:t")):
                    text_parts.append(t_elem.text or "")
            if base_rPr is None:
                rPr = child.find(qn("w:rPr"))
                if rPr is not None:
                    base_rPr = deepcopy(rPr)

        elif child.tag == qn("w:hyperlink"):
            contador += 1
            r_id = child.get(qn("r:id"), "")
            anchor = child.get(qn("w:anchor"), "")
            url = ""
            if r_id:
                try:
                    url = parrafo.part.rels[r_id].target_ref
                except (KeyError, AttributeError):
                    pass

            link_text = ""
            link_rPr = None
            for run_elem in child.findall(qn("w:r")):
                for t_elem in run_elem.findall(qn("w:t")):
                    link_text += t_elem.text or ""
                if link_rPr is None:
                    found = run_elem.find(qn("w:rPr"))
                    if found is not None:
                        link_rPr = deepcopy(found)

            link_map[contador] = HyperlinkInfo(
                url=url, r_id=r_id, anchor=anchor, rPr=link_rPr, texto=link_text,
            )
            text_parts.append(_token(contador))

    return "".join(text_parts).strip(), link_map, nota_map, base_rPr


def _extraer_de_parrafos(parrafos: list[Paragraph]) -> list[UnidadTraducible]:
    """Extrae unidades traducibles de una lista de párrafos."""
    unidades = []
    for parrafo in parrafos:
        if _parrafo_tiene_imagen(parrafo):
            continue
        texto, links, notas, base_rPr = _extraer_texto_con_links(parrafo)
        if not texto:
            continue
        unidades.append(UnidadTraducible(
            texto=texto, parrafo=parrafo,
            hyperlinks=links, notas=notas, base_rPr=base_rPr,
        ))
    return unidades


def _recorrer_unidades(parrafos, tablas, unidades: list) -> None:
    """Recorre recursivamente párrafos y tablas (incluyendo anidadas)
    recolectando unidades de texto traducibles."""
    unidades.extend(_extraer_de_parrafos(parrafos))
    for tabla in tablas:
        for fila in tabla.rows:
            for celda in fila.cells:
                _recorrer_unidades(celda.paragraphs, celda.tables, unidades)


# Puntuación que indica fin de oración/párrafo completo
_CIERRE_PATTERN = re.compile(r'[.!?:…]$|\.{3,}$')


def _mergear_parrafos_xml(p1: Paragraph, p2: Paragraph) -> None:
    """Mueve los runs de p2 al final de p1 y elimina p2 del documento."""
    # Insertar un espacio entre los textos mergeados
    space_run = OxmlElement("w:r")
    space_t = OxmlElement("w:t")
    space_t.text = " "
    space_t.set(qn("xml:space"), "preserve")
    space_run.append(space_t)
    p1._element.append(space_run)

    # Mover todos los runs de p2 a p1
    for child in list(p2._element):
        p1._element.append(child)

    # Eliminar p2 del documento
    parent = p2._element.getparent()
    if parent is not None:
        parent.remove(p2._element)


def mergear_parrafos(unidades: list[UnidadTraducible]) -> list[UnidadTraducible]:
    """Mergea párrafos consecutivos partidos por soft-breaks (PDF→DOCX Calibre).

    Dos párrafos se mergean si el primero no termina con puntuación de cierre
    (. ! ? : … ...) y el segundo arranca con minúscula (continuación de frase).
    La operación modifica el XML del documento: mueve los runs del segundo
    párrafo al primero y elimina el segundo.
    """
    if not unidades:
        return []

    resultado = []
    actual = unidades[0]

    for siguiente in unidades[1:]:
        t1 = actual.texto.rstrip()
        t2 = siguiente.texto.lstrip()
        termina_cierre = bool(_CIERRE_PATTERN.search(t1))
        empieza_minus = t2 and t2[0].islower()

        if not termina_cierre and empieza_minus:
            actual.texto = t1 + " " + t2
            _mergear_parrafos_xml(actual.parrafo, siguiente.parrafo)
        else:
            resultado.append(actual)
            actual = siguiente

    resultado.append(actual)
    return resultado


def _formato_signature(run) -> str:
    """Devuelve el XML del w:rPr del run como string para comparar formato."""
    rPr = run._element.find(qn("w:rPr"))
    return tostring(rPr, encoding="unicode") if rPr is not None else ""


def desagrupar_formatos(unidades: list[UnidadTraducible]) -> list[UnidadTraducible]:
    """Parte párrafos con formato mixto (negrita + normal) en unidades separadas.

    Cada grupo de runs consecutivos con el mismo formato se traduce en su propia
    llamada, preservando el formato sin heurística de split sobre la traducción.
    Párrafos con hipervínculos o notas se dejan intactos.
    """
    resultado = []
    for unidad in unidades:
        if unidad.hyperlinks or unidad.notas:
            resultado.append(unidad)
            continue

        runs = unidad.parrafo.runs
        if len(runs) <= 1:
            resultado.append(unidad)
            continue

        # Agrupar runs consecutivos comparando el XML de w:rPr
        grupos = []  # [(signature, [run_indices])]
        for i, run in enumerate(runs):
            texto_run = (run.text or "").strip()
            if not texto_run:
                continue  # runs de solo whitespace se ignoran
            sig = _formato_signature(run)
            if grupos and grupos[-1][0] == sig:
                grupos[-1][1].append(i)
            else:
                grupos.append((sig, [i]))

        if len(grupos) <= 1:
            resultado.append(unidad)
            continue

        # Crear una unidad por grupo de runs con mismo formato
        for _sig, indices in grupos:
            texto = " ".join((runs[i].text or "").strip() for i in indices).strip()
            if not texto:
                continue
            resultado.append(UnidadTraducible(
                texto=texto,
                parrafo=unidad.parrafo,
                base_rPr=unidad.base_rPr,
                run_indices=indices,
            ))

    return resultado


def extraer_unidades(ruta_docx: Path) -> tuple[Document, list[UnidadTraducible]]:
    """Abre un DOCX y extrae todas las unidades traducibles.
    Recursivo: cuerpo principal, tablas, y tablas anidadas dentro de celdas.
    Devuelve (documento, lista_de_unidades).
    """
    doc = Document(str(ruta_docx))
    unidades = []
    _recorrer_unidades(doc.paragraphs, doc.tables, unidades)
    return doc, unidades


def _recorrer_imagenes(parrafos, tablas, imagenes: list) -> None:
    """Recorre recursivamente párrafos y tablas (incluyendo anidadas) recolectando imágenes."""
    for parrafo in parrafos:
        imagenes.extend(_extraer_imagenes_de_parrafo(parrafo))
    for tabla in tablas:
        for fila in tabla.rows:
            for celda in fila.cells:
                _recorrer_imagenes(celda.paragraphs, celda.tables, imagenes)


def extraer_imagenes(doc: Document) -> list[ImagenTraducible]:
    """Recorre el documento y devuelve todas las imágenes inline traducibles.
    Recursivo: cuerpo principal, tablas, y tablas anidadas dentro de celdas.
    """
    imagenes = []
    _recorrer_imagenes(doc.paragraphs, doc.tables, imagenes)
    return imagenes


# Token opaco que representa un hipervínculo o nota. Los corchetes matemáticos
# ⟦ ⟧ no aparecen en texto real y el modelo los arrastra sin tocarlos
# porque no tiene que traducir nada adentro (solo un número).
_TOKEN_PATTERN = re.compile(r"⟦(\d+)⟧")


def _token(idx: int) -> str:
    return f"⟦{idx}⟧"


def _crear_run_element(texto: str, rPr: object = None) -> OxmlElement:
    """Crea un elemento w:r con texto y formato opcional."""
    run_el = OxmlElement("w:r")
    if rPr is not None:
        run_el.append(deepcopy(rPr))
    t_el = OxmlElement("w:t")
    t_el.text = texto
    t_el.set(qn("xml:space"), "preserve")
    run_el.append(t_el)
    return run_el


def _crear_run_nota(info: NotaInfo) -> OxmlElement:
    """Crea un elemento w:r con referencia a nota al pie o final."""
    run_el = OxmlElement("w:r")
    if info.rPr is not None:
        run_el.append(deepcopy(info.rPr))
    tag = "w:footnoteReference" if info.tipo == "footnote" else "w:endnoteReference"
    ref_el = OxmlElement(tag)
    ref_el.set(qn("w:id"), info.nota_id)
    run_el.append(ref_el)
    return run_el


def _aplicar_estilo_link(run_el: OxmlElement):
    """Aplica color azul y subrayado a un run de hipervínculo."""
    rPr = run_el.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        run_el.insert(0, rPr)

    # Color azul estándar de hyperlinks
    color_el = rPr.find(qn("w:color"))
    if color_el is None:
        color_el = OxmlElement("w:color")
        rPr.append(color_el)
    color_el.set(qn("w:val"), "0000FF")

    # Subrayado simple
    u_el = rPr.find(qn("w:u"))
    if u_el is None:
        u_el = OxmlElement("w:u")
        rPr.append(u_el)
    u_el.set(qn("w:val"), "single")


def _reconstruir_parrafo(parrafo: Paragraph, texto: str,
                         link_map: dict, nota_map: dict,
                         base_rPr: object):
    """Reconstruye el XML del párrafo con hipervínculos y notas preservados."""
    elem = parrafo._element

    # Limpiar todos los hijos excepto w:pPr (propiedades del párrafo)
    for child in list(elem):
        if child.tag != qn("w:pPr"):
            elem.remove(child)

    # Procesar texto: partir por tokens ⟦N⟧ y reconstruir secuencialmente
    last_end = 0
    for match in _TOKEN_PATTERN.finditer(texto):
        # Texto plano antes del token
        plain = texto[last_end:match.start()]
        if plain:
            elem.append(_crear_run_element(plain, base_rPr))
        last_end = match.end()

        idx = int(match.group(1))

        # ¿Es un link? Se reinserta con su texto visible ORIGINAL (no traducido).
        if idx in link_map:
            info = link_map[idx]
            hl_el = OxmlElement("w:hyperlink")
            if info.r_id:
                hl_el.set(qn("r:id"), info.r_id)
            if info.anchor:
                hl_el.set(qn("w:anchor"), info.anchor)
            link_rPr = info.rPr if info.rPr is not None else base_rPr
            link_run = _crear_run_element(info.texto, link_rPr)
            _aplicar_estilo_link(link_run)
            hl_el.append(link_run)
            elem.append(hl_el)
        # ¿Es una nota?
        elif idx in nota_map:
            elem.append(_crear_run_nota(nota_map[idx]))
        # Token desconocido: se ignora (no debería pasar; la validación previa
        # en aplicar_traducciones evita reconstruir párrafos con tokens rotos).

    # Texto plano restante después del último token
    if last_end < len(texto):
        remaining = texto[last_end:]
        if remaining:
            elem.append(_crear_run_element(remaining, base_rPr))


def _tokens_intactos(unidad: UnidadTraducible) -> bool:
    """True si la traducción conserva exactamente los tokens ⟦N⟧ esperados
    (cada uno una vez, sin faltantes ni sobrantes)."""
    esperados = set(unidad.hyperlinks) | set(unidad.notas)
    encontrados = [int(n) for n in _TOKEN_PATTERN.findall(unidad.traduccion)]
    return sorted(encontrados) == sorted(esperados)


def aplicar_traducciones(unidades: list[UnidadTraducible]) -> int:
    """Reescribe los párrafos con las traducciones, preservando hipervínculos y notas.

    Devuelve la cantidad de párrafos que quedaron sin traducir porque el modelo
    rompió los tokens ⟦N⟧ (fallback: se deja el original con sus links intactos,
    en vez de escupir basura o perder referencias).
    """
    fallbacks = 0
    for unidad in unidades:
        if not unidad.traduccion:
            continue

        parrafo = unidad.parrafo

        if unidad.hyperlinks or unidad.notas:
            # Red de seguridad: si el modelo rompió los tokens, no reconstruimos
            # (dejaríamos links perdidos o marcadores crudos visibles). Preferimos
            # dejar el párrafo original, con sus hipervínculos y notas intactos.
            if not _tokens_intactos(unidad):
                fallbacks += 1
                continue
            _reconstruir_parrafo(
                parrafo, unidad.traduccion,
                unidad.hyperlinks, unidad.notas, unidad.base_rPr,
            )
        else:
            # Sin hipervínculos ni notas.
            # Limpiar tokens ⟦N⟧ que el modelo pueda haber desplazado hasta acá desde
            # un párrafo vecino: en un párrafo sin links no significan nada y quedarían
            # como basura visible.
            texto = _TOKEN_PATTERN.sub("", unidad.traduccion)
            texto = re.sub(r"  +", " ", texto).strip()
            runs = parrafo.runs
            if not runs:
                continue

            if unidad.run_indices is not None:
                # Esta unidad solo modifica los runs que le pertenecen
                for i, idx in enumerate(unidad.run_indices):
                    if i == 0:
                        runs[idx].text = texto
                    else:
                        runs[idx].text = ""
            else:
                # Párrafo homogéneo: todo al primer run (hereda su formato)
                runs[0].text = texto
                for run in runs[1:]:
                    run.text = ""

    return fallbacks


def _aplicar_fuente_parrafos(parrafos: list[Paragraph], nombre_fuente: str,
                             tamano: int):
    """Aplica fuente y tamaño a todos los runs de una lista de párrafos."""
    for parrafo in parrafos:
        for run in parrafo.runs:
            run.font.name = nombre_fuente
            run.font.size = Pt(tamano)


def _aplicar_fuente_recursivo(parrafos, tablas, nombre_fuente: str, tamano: int):
    """Aplica fuente recursivamente a párrafos y tablas (incluyendo anidadas)."""
    _aplicar_fuente_parrafos(parrafos, nombre_fuente, tamano)
    for tabla in tablas:
        for fila in tabla.rows:
            for celda in fila.cells:
                _aplicar_fuente_recursivo(celda.paragraphs, celda.tables,
                                          nombre_fuente, tamano)


def aplicar_fuente(doc: Document, nombre_fuente: str, tamano: int):
    """Aplica fuente y tamaño a todo el documento (cuerpo + tablas, incluyendo anidadas)."""
    _aplicar_fuente_recursivo(doc.paragraphs, doc.tables, nombre_fuente, tamano)


def guardar_docx(doc: Document, ruta_salida: Path):
    """Guarda el documento DOCX."""
    doc.save(str(ruta_salida))


# --- Aplicación de captions traducidos a imágenes ---


def _crear_parrafo_caption(texto: str) -> OxmlElement:
    """Crea un párrafo en cursiva, tamaño 90%, centrado, con el caption."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    pPr.append(jc)
    # rPr base del párrafo (afecta marcador de fin de párrafo)
    base_rPr = OxmlElement("w:rPr")
    i_el = OxmlElement("w:i")
    base_rPr.append(i_el)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")  # 10pt = 20 half-points
    base_rPr.append(sz)
    pPr.append(base_rPr)
    p.append(pPr)

    # Run con el texto
    r = OxmlElement("w:r")
    r_rPr = OxmlElement("w:rPr")
    i2 = OxmlElement("w:i")
    r_rPr.append(i2)
    sz2 = OxmlElement("w:sz")
    sz2.set(qn("w:val"), "20")
    r_rPr.append(sz2)
    r.append(r_rPr)
    t = OxmlElement("w:t")
    t.text = texto
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)
    return p


def _parrafo_solo_tiene_runs_vacios(elem_p) -> bool:
    """True si el párrafo no contiene runs con texto ni más imágenes."""
    for r in elem_p.findall(qn("w:r")):
        # ¿texto?
        for t in r.findall(qn("w:t")):
            if (t.text or "").strip():
                return False
        # ¿imagen u otra cosa visual?
        if r.find(qn("w:drawing")) is not None or r.find(qn("w:pict")) is not None:
            return False
    return True


def aplicar_captions_imagenes(imagenes: list[ImagenTraducible]) -> int:
    """Para cada imagen con traducción, inserta un párrafo con el texto traducido
    debajo de la imagen. Sin tabla.

    Devuelve la cantidad de captions aplicados.
    """
    aplicados = 0
    for img in imagenes:
        if not img.traduccion:
            continue
        elem_p = img.parrafo._element
        parent = elem_p.getparent()
        if parent is None:
            continue

        # Verificar que el run de la imagen sigue dentro del párrafo
        if img.run_element.getparent() is not elem_p:
            continue

        # Crear párrafo con el texto traducido como caption debajo de la imagen
        cap = _crear_parrafo_caption(img.traduccion)
        # Insertarlo justo después del párrafo que contiene la imagen
        elem_p.addnext(cap)

        aplicados += 1
    return aplicados
